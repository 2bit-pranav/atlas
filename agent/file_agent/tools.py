import csv
from pathlib import Path
from typing import Any, Dict, List, Optional
import docx
import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
import subprocess
import sys
import shutil
import tempfile

DOWNLOADS_DIR = Path.home() / "Downloads"

def _get_target_path(file_name: str) -> Path:
    """Resolves target path in Downloads, auto-incrementing if the filename already exists."""
    DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)
    clean_name = Path(file_name).name
    target = DOWNLOADS_DIR / clean_name
    counter = 1
    while target.exists():
        target = DOWNLOADS_DIR / f"{target.stem}_copy_{counter}{target.suffix}"
        counter += 1
    return target

def read_file(file_path: str) -> str:
    """Reads a UTF-8 text file from an arbitrary path on disk."""
    path = Path(file_path).expanduser().resolve()
    if not path.exists() or not path.is_file():
        return f"[FILE_ERROR]: File '{file_path}' does not exist or is not a file."
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
        return f"SUCCESS: Read '{path.name}'.\n--- CONTENT ---\n{content}"
    except Exception as e:
        return f"[FILE_ERROR]: Failed to read file '{file_path}': {e}"

def save_text_file(file_name: str, content: str) -> str:
    """Saves text, Markdown, or CSV data directly into the Downloads folder."""
    try:
        target = _get_target_path(file_name)
        target.write_text(content, encoding="utf-8")
        return f"SUCCESS: Saved file to '{target.resolve()}' ({target.stat().st_size} bytes)."
    except Exception as e:
        return f"[FILE_ERROR]: Failed to save text file '{file_name}': {e}"

def create_pdf_document(file_name: str, title: str, sections: List[Dict[str, str]]) -> str:
    """Generates a structured PDF document in the Downloads folder."""
    try:
        target = _get_target_path(file_name)
        doc = SimpleDocTemplate(str(target), pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph(title, styles["Title"]), Spacer(1, 14)]
        for sec in sections:
            if heading := sec.get("heading"):
                story.append(Paragraph(heading, styles["Heading2"]))
                story.append(Spacer(1, 6))
            if body := sec.get("body"):
                formatted_body = body.replace("\n", "<br/>")
                story.append(Paragraph(formatted_body, styles["Normal"]))
                story.append(Spacer(1, 12))
        doc.build(story)
        return f"SUCCESS: Generated PDF at '{target.resolve()}' ({target.stat().st_size} bytes)."
    except Exception as e:
        return f"[FILE_ERROR]: Failed to create PDF '{file_name}': {e}"

def create_docx_document(file_name: str, title: str, sections: List[Dict[str, str]]) -> str:
    """Generates a Microsoft Word (.docx) document in the Downloads folder."""
    try:
        target = _get_target_path(file_name)
        doc = docx.Document()
        doc.add_heading(title, level=0)
        for sec in sections:
            if heading := sec.get("heading"):
                doc.add_heading(heading, level=1)
            if body := sec.get("body"):
                doc.add_paragraph(body)
        doc.save(str(target))
        return f"SUCCESS: Generated DOCX at '{target.resolve()}' ({target.stat().st_size} bytes)."
    except Exception as e:
        return f"[FILE_ERROR]: Failed to create DOCX '{file_name}': {e}"

def create_excel_spreadsheet(
    file_name: str, headers: List[str], rows: List[List[Any]], sheet_name: str = "Sheet1"
) -> str:
    """Generates an Excel (.xlsx) workbook in the Downloads folder."""
    try:
        target = _get_target_path(file_name)
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name
        ws.append(headers)
        header_fill = PatternFill(start_color="1F497D", end_color="1F497D", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")
        for row in rows:
            ws.append(row)
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            col_letter = openpyxl.utils.get_column_letter(col[0].column)
            ws.column_dimensions[col_letter].width = max(max_len + 3, 10)
        wb.save(str(target))
        return f"SUCCESS: Generated Excel sheet at '{target.resolve()}' ({target.stat().st_size} bytes)."
    except Exception as e:
        return f"[FILE_ERROR]: Failed to create Excel file '{file_name}': {e}"

def verify_file(file_path: str) -> str:
    """Validates that a created file exists and contains non-zero byte size."""
    path = Path(file_path).expanduser().resolve()
    if not path.exists() or not path.is_file():
        return f"VERIFY_FAILED: File '{file_path}' does not exist or is invalid."
    size = path.stat().st_size
    if size == 0:
        return f"VERIFY_FAILED: File '{file_path}' is empty (0 bytes)."
    return f"VERIFY_SUCCESS: Verified '{path.name}' at '{path.parent}' ({size} bytes)."

def run_python_code(
    code: str,
    script_name: str = "temp_skill_script.py",
    dependencies: Optional[List[str]] = None,
) -> str:
    """
    Executes Python code in an isolated workspace.
    If dependencies are passed, provisions them dynamically using 'uv run'.
    """
    try:
        work_dir = DOWNLOADS_DIR
        work_dir.mkdir(parents=True, exist_ok=True)
        script_path = work_dir / script_name
        script_path.write_text(code, encoding="utf-8")

        # Build command: use 'uv run' if dependencies are specified or uv is available
        uv_path = shutil.which("uv")
        if dependencies and uv_path:
            cmd = [uv_path, "run"]
            for dep in dependencies:
                cmd.extend(["--with", dep])
            cmd.append(str(script_path))
        else:
            cmd = [sys.executable, str(script_path)]

        res = subprocess.run(
            cmd,
            cwd=str(work_dir),
            capture_output=True,
            text=True,
            timeout=120,
        )

        if res.returncode == 0:
            return f"[PYTHON_EXEC_SUCCESS]\nSTDOUT:\n{res.stdout}"

        stderr_text = res.stderr or res.stdout
        return f"[PYTHON_RUNTIME_ERROR] (Exit Code {res.returncode}):\nSTDERR:\n{stderr_text}"

    except subprocess.TimeoutExpired:
        return "[EXECUTION_TIMEOUT]: Script execution timed out after 120 seconds."
    except Exception as e:
        return f"[PYTHON_EXEC_ERROR]: Failed to execute Python script: {e}"

def run_terminal_command(command: str) -> str:
    """Executes a shell command inside the host execution workspace with error feedback."""
    try:
        res = subprocess.run(
            command,
            shell=True,
            cwd=str(DOWNLOADS_DIR),
            capture_output=True,
            text=True,
            timeout=120,
        )
        if res.returncode == 0:
            return f"[TERMINAL_EXEC_SUCCESS]\nSTDOUT:\n{res.stdout}"
        return f"[TERMINAL_EXEC_ERROR] (Exit Code {res.returncode}):\nSTDERR:\n{res.stderr}\nSTDOUT:\n{res.stdout}"
    except subprocess.TimeoutExpired:
        return "[EXECUTION_TIMEOUT]: Shell command execution timed out after 120 seconds."
    except Exception as e:
        return f"[TERMINAL_EXEC_ERROR]: Failed to execute command '{command}': {e}"